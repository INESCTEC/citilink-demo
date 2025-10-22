import os
from docx import Document
import subprocess
import re
from docx.table import Table, _Cell
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import argparse
from typing import Tuple
import os
import subprocess
from pathlib import Path
import re
from datetime import datetime, date



def date_of_header(header: str) -> date:
    """
    Recebe um header e retorna a data mais recente que ele contém.
    Suporta dd/mm/yyyy, dd-mm-yyyy, dd.mm.yyyy com ano de 2 ou 4 dígitos.
    """
    date_formats:list[str] = ['%d/%m/%Y', '%d-%m-%Y', '%d.%m.%Y',
                    '%d/%m/%y', '%d-%m-%y', '%d.%m.%y']
    
    date_regex:str = r'(\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})'
    matches = re.findall(date_regex, header)

    parsed_dates = []
    for match in matches:
        for fmt in date_formats:
            try:
                parsed_date = datetime.strptime(match, fmt).date()
                parsed_dates.append(parsed_date)
                break  
            except ValueError:
                continue
    return max(parsed_dates) if parsed_dates else None


def the_newest_header(header_list: list[str]) -> str:
    """
    Recebe uma lista de headers, e retorna aquele que tiver a data mais recente.
    """
    date_list = []

    for header in header_list:
        date = date_of_header(header)  
        date_list.append((header, date))

    newest_date = date_list[0]

    for i in range(1, len(date_list)):
        if date_list[i][1] > newest_date[1]:
            newest_date = date_list[i]

    return newest_date[0] if len(newest_date) > 0 else ""

    
    

def docx_to_txt_pandoc(input_path: str, output_path: str, _extract_header: bool = False) -> None:
    input_file = Path(input_path)
    output_file = Path(output_path)

    if not input_file.exists():
        raise FileNotFoundError(f"Input file not found: {input_file}")

    # Converte o DOCX para plain text com pandoc
    result = subprocess.run(
        ['pandoc', str(input_file), '-t', 'plain+smart'],
        capture_output=True,
        text=True
    )

    if result.returncode != 0:
        print("Pandoc error:", result.stderr)
        return

    # Junta linhas dentro de parágrafos; separa parágrafos reais

    header = ""

    if _extract_header:
        header = extract_header(input_path)
    
        header = the_newest_header(header)


    lines = result.stdout.splitlines()
    paragraphs = []
    current_paragraph = []

    for line in lines:
        stripped = line.strip()

        if stripped == "":
            # Quando há uma linha em branco, fecha o parágrafo atual
            if current_paragraph:
                paragraph_text = " ".join(current_paragraph).strip()
                if paragraph_text:
                    paragraphs.append(paragraph_text)
                current_paragraph = []
        else:
            current_paragraph.append(stripped)

    # Adiciona o último parágrafo, se existir
    if current_paragraph:
        paragraph_text = " ".join(current_paragraph).strip()
        if paragraph_text:
            paragraphs.append(paragraph_text)

    # Escreve os parágrafos separados por uma quebra de linha (sem linhas em branco extras)
    cleaned_strings = [re.sub(r'[-]{3,}|[.]{3,}', '', s) for s in paragraphs]


    cleaned_strings.insert(0, header)

    # print("Header:", header)
    # print("Paragraphs:", cleaned_strings)

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("\n".join(cleaned_strings))  # uma quebra de linha entre parágrafos

    # print(f"Texto limpo salvo em {output_file}")


def docx_of_doc(input_file: str) -> str:
    file = Path(input_file)
    if not file.exists():
        raise FileNotFoundError(f"Ficheiro {file} não encontrado.")
    
    if file.suffix.lower() != '.doc':
        raise ValueError("O ficheiro de entrada tem que ser um .doc")
    
    docx_file = file.with_suffix(".docx")
    try:
        subprocess.run(["soffice", "--headless", "--convert-to", "docx", str(file)], check=True)
        # print(f"Convertido: {file} -> {docx_file}")
        os.remove(file)
        # print(f"Ficheiro removido: {file}")
        return str(docx_file)
    except subprocess.CalledProcessError as e:
        # print(f"Erro ao converter {file}: {e}")
        raise
 
def docx_of_pdf(input_file: str) -> str:
    """Converte ficheiro .pdf -> .docx """
    file = Path(input_file)
    if not file.exists():
        raise FileNotFoundError(f"Ficheiro {file} não encontrado.")
    
    if file.suffix.lower() != '.pdf':
        raise ValueError("O ficheiro de input deve ser um .pdf")
    
    docx_file = file.with_suffix(".docx")
    try:
        #subprocess.run(["soffice", "--headless", "--convert-to", "docx", str(file)], check=True)
        subprocess.run(["libreoffice", "--headless", "--convert-to", "html", str(file)], check=True)

        subprocess.run(["soffice", "--convert-to", "docx:MS Word 2007 XML", str(file.with_suffix(".html"))], check=True)
        # print(f"Convertido: {file} -> {docx_file}")
        return str(docx_file)
    except subprocess.CalledProcessError as e:
        print(f"Erro ao converter {file}: {e}")
        raise
 

def extract_header(docx_path):
    doc = Document(docx_path)
    headers = []

    end=False
    
    for section in doc.sections:
        for header_part in [section.header, section.first_page_header, section.even_page_header]:
            if header_part is not None:
                header_text = set()
                for paragraph in header_part.paragraphs:
                    if paragraph.text.strip():
                        header_text.add(paragraph.text)
                for table in header_part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip() and not end:
                                #print(cell.text)
                                header_text.add(cell.text)
                                end=True
                
                if header_text:
                    headers.append("\n".join(list(header_text)))
    
    
    return headers
 
def docx_of_doc(doc_path:str) -> str:
    docx_path:str = doc_path + "x"
    path, filename = os.path.split(doc_path)
 
    process:subprocess.CompletedProcess=subprocess.run(["soffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", path], check=True)
    return docx_path if process.returncode == 0 else ""
 
 
def txt_of_docx(docx_path:str, txt_path:str, title_of_header:bool, entidade:str, no_caption:bool=False) -> None:
    doc:Document = Document(docx_path)
    previous_tex:str=""
    previous_align=None
    title:str=""
    
    if len(entidade) > 0:
        title=entidade+".\n"

         
         
    

    title += the_newest_header(list(set(extract_header(docx_path))))  if title_of_header else ""
 
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        if title_of_header:
            txt_file.write(title+'\n')
        for element in doc.element.body:
            if element.tag.endswith('p'):
                text:str = (element.text).strip()
                if len(text) > 0:
                    if len(previous_tex) > 0:
                        _:int=txt_file.write(previous_tex + '\n')
 
                    if element.alignment!=WD_PARAGRAPH_ALIGNMENT.CENTER:
                        _:int=txt_file.write(text + '\n')
                        previous_tex=""
                    else:
                        previous_tex=text
                        previous_align=element.alignment
                    
            elif element.tag.endswith('tbl'):  
                _:int=txt_file.write('\n')
                table:Table = next(t for t in doc.tables if t._element == element)
                csv_lines:list[str] = []
                #if previous_align==WD_PARAGRAPH_ALIGNMENT.CENTER:
                #    print(previous_tex)
                    
                for row in table.rows:
                    row_data=[cell.text.strip().replace('\n', '') for cell in row.cells]
                    if all(cell != '' for cell in row_data):  # Only add row if no empty cells
                            csv_lines.append(",".join(row_data))

                _:int=txt_file.write(".init_table\n")
                length_of_header:int=0
                if previous_align==WD_PARAGRAPH_ALIGNMENT.CENTER and (not no_caption):
                    header:str=f"————————————{previous_tex}————————————\n"
                    length_of_header=len(header)
                    _:int=txt_file.write(header)
                else:
                    header:str="———————————— TABELA ————————————\n" if not no_caption else "——————————————————————————\n"
                    length_of_header=len(header)
                    _:int=txt_file.write(header)
                _:int=txt_file.write("\n".join(csv_lines) + '\n')
                _:int=txt_file.write(length_of_header*'—'+'\n')
                _:int=txt_file.write(".end_table\n\n")
                previous_tex = ""
 
def remove_spaces(text:str) -> str:
    return re.sub(r'-{3,}', '', re.sub(r' {2,}', ' ', text))
 
def string_of_txt(file_path:str) -> str:
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()  
    
def correct_broke_phrases(text:str) -> str:
    result:list[str] = []
    prev_char = None
    inside_table:bool = False
 
    i:int = 0
    while i < len(text):
        
        if ".init_table" in text[i:i+12] :
            #print("GOT IT")
            inside_table = True
            #result.append(text[i:i+12])
            i += 12
        # Detecta fim da tabela
        elif ".end_table" in text[i:i+11] :
            #print("BYE")
            inside_table = False
            #result.append(text[i:i+11])
            i += 11
        # Processa texto normal (fora das tabelas)
        elif not inside_table:
            
            if text[i] == "\n":
                if i+3 < len(text) and text[i+1].isalpha() and text[i+2] ==')' and text[i + 3] == ' ':
                    result.append("\n")
                if i+3 < len(text) and prev_char == "." and text[i + 1] == ' ' and text[i + 2] == ' ' and text[i+3].isupper() :
                    #print("INNNN")
                    result.append("\n")
                elif i+1 < len(text) and text[i+1] == "-":
                    result.append("\n")
                elif prev_char == ":":
                    result.append("\n")
                elif (prev_char is not None) and prev_char != "." and (not prev_char.isupper() and prev_char != " "):
                    result.append(" ")
                elif prev_char == ";":
                    result.append("\n")
                else:
                    result.append("\n")
            else:
                result.append(text[i])  
            prev_char = text[i]  
            i += 1
        else:
            #print(inside_table)
            if text[i] == "\n":
                result.append("\n")
            else:
                result.append(text[i])
            i += 1
 
    return ''.join(result)
 
#---------------------------------------------------------------------------
def standard_txt_of_docx(fin: str, fout: str, title_of_header:bool=False, entidade:str="", no_caption:bool=False) -> int:
    _docx_path:str = fin
    if fin.endswith(".doc"):
        _docx_path = docx_of_doc(fin)
    if len(_docx_path) > 0:
 
        txt_of_docx(_docx_path, fout, title_of_header, entidade,no_caption=no_caption)

        content:str=string_of_txt(fout)
        # print(content)
        content=remove_spaces(content)
        content = correct_broke_phrases(content)
 
        with open(fout, 'w', encoding='utf-8') as file:

            file.write(content)
        # print(f"Conversão completa: {fout}")
    else:
        print("Conversão falhou.")
 
 
 
#------------- covilha ------------------
 
 
def copy_table(new_doc, table):
    new_table = new_doc.add_table(rows=0, cols=len(table.columns))
    new_table.style = table.style
 
    # Copiar células da tabela
    for row in table.rows:
        row_cells = new_table.add_row().cells
        for i, cell in enumerate(row.cells):
            row_cells[i].text = cell.text.strip()
 
    return new_table
 
def process_cell_content(cell):
    content:list[Tuple[str, str]] = []
    
    for elem in cell._element:
        if elem.tag.endswith('p'):  # Parágrafos
            para_text = ''.join(node.text for node in elem.iter() if node.tag.endswith('t') and node.text).strip()
            if para_text:
                content.append(('p', para_text))
        elif elem.tag.endswith('tbl'):  # Tabelas
            for nested_table in cell.tables:
                if nested_table._element == elem:
                    content.append(('t', nested_table))
                    break
 
    return content
 
def create_new_document(input_path:str, output_path:str) -> int:
    doc:Document = Document(input_path)

    
    table:Table = doc.tables[0]
    cell:_Cell = table.cell(0, 0)
 
    new_doc:Document = Document()
 
    for elem_type, elem in process_cell_content(cell):
        #print(elem)
        if elem_type == 'p':
            new_doc.add_paragraph(elem)
        elif elem_type == 't':
            copy_table(new_doc, elem)
 
    try:
        new_doc.save(output_path)
        return 0
    except:
        return 1
 
 
def CM_campo_maior(fin:str, fout:str) -> int:
    return standard_txt_of_docx(fin, fout)
 
def CM_alandroal(fin:str, fout:str, entidade:str="") -> int:
    return standard_txt_of_docx(fin, fout, 
                                title_of_header=True, 
                                entidade=entidade)
 
 
 
 
def preprocessing_covilha(fin:str, fout:str) -> int:
    _exec_code:int = create_new_document(fin, "aux.docx")
    return _exec_code
 
 
def CM_Covilha(fin:str, fout:str) -> int:
    _=preprocessing_covilha(fin, fin)
    _exec_code:int = standard_txt_of_docx("aux.docx", fout, no_caption=True)
    os.remove("aux.docx")
 
    return _exec_code
 
 
def main(fin:str, fout:str, entidade:str) -> int:

    if fin.endswith(".pdf"):
        fin = docx_of_pdf(fin)
    elif fin.endswith(".doc"):
        fin = docx_of_doc(fin)
    
    if entidade.startswith("CM"):
 
        if "COVILHA" in entidade:
            CM_Covilha(fin, fout)
        elif "ALANDROAL" in entidade:
            docx_to_txt_pandoc(fin, fout, _extract_header=True)
        elif "CAMPOMAIOR" in entidade:
            CM_campo_maior(fin, fout)
        elif "FUNDAO" in entidade:
            CM_alandroal(fin, fout)
        elif "PORTO" in entidade:
            docx_to_txt_pandoc(fin, fout)
        elif "GUIMARAES" in entidade:
            CM_alandroal(fin, fout, 
                         entidade="Câmara Municipal de Guimarães")
        else:
            print("Sem suporte") 
    
    elif entidade.startswith("AM"):
        if "PORTO" in entidade:
            docx_to_txt_pandoc(fin, fout)
        elif "FUNDAO" in entidade:
            CM_alandroal(fin, fout)
        elif "ALANDROAL" in entidade: 
            CM_alandroal(fin, fout)
        elif "GUIMARAES" in entidade:
            CM_alandroal(fin, fout, 
                         entidade="Assembleia Municipal de Guimarães")
        elif "COVILHA" in entidade:
            CM_Covilha(fin, fout)
        elif "CAMPOMAIOR" in entidade:
            CM_campo_maior(fin, fout)
        else:
            print("Sem suporte") 
 
 
 
 
parser:argparse.ArgumentParser = argparse.ArgumentParser(prog='python txt_of_doc.py')
parser.add_argument('--input', help='caminho do ficheiro de input (.doc|.docx|.pdf)', type=str, required=True)
parser.add_argument('--output', help='caminho onde o ficheiro .txt é guardado', default=subprocess.STDOUT)
parser.add_argument('--entidade', help='entidade a qual pertence o input', type=str, required=True)
 
"""
{CM|AM}_COVILHA
{CM|AM}_ALANDROAL
{CM|AM}_CAMPOMAIOR
{CM|AM}_PORTO
{CM|AM}_GUIMARAES
"""
 
# Only execute this code if the script is run directly, not when imported
if __name__ == "__main__":
    args = parser.parse_args()
    docx_fpath=args.input
    txt_fpath=args.output
    entidade = args.entidade
    main(docx_fpath, txt_fpath, entidade)